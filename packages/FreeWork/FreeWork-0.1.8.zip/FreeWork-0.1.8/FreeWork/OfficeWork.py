import openpyxl
import docx
import shutil
import os
import fiona
import pandas as pd
from docx.shared import Cm
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def FileName(FilePath):
    return os.path.basename(FilePath)


def FileFatherPath(FilePath):
    return os.path.dirname(FilePath)


def CopyFile(FileOriginalPath, FileNewPath):
    FileNewFatherPath = os.path.dirname(FileNewPath)
    if not os.path.exists(FileNewFatherPath):
        os.makedirs(FileNewFatherPath)
    shutil.copy2(FileOriginalPath, FileNewPath)


def MoveFile(FileOriginalPath, FileNewPath):
    FileNewFatherPath = os.path.dirname(FileNewPath)
    if not os.path.exists(FileNewFatherPath):
        os.makedirs(FileNewFatherPath)
    shutil.move(FileOriginalPath, FileNewPath)


def ExcleRead(ExclePath, SheetIndex, Rowlow, Rowmax, Collow, Colmax):
    if Rowlow > Rowmax:
        t = Rowmax
        Rowmax = Rowlow
        Rowlow = t
    if Collow > Colmax:
        t = Rowmax
        Colmax = Collow
        Collow = t
    RowNum = Rowmax - Rowlow + 1
    ColNum = Colmax - Collow + 1
    # 打开工作簿
    workbook = openpyxl.load_workbook(ExclePath)
    # 获取所有工作表
    sheets = workbook.sheetnames
    # 选择第一个工作表
    sheet = workbook[sheets[SheetIndex - 1]]
    # 存储为列表

    m = 0
    i = 0
    SheetData = [[None for j in range(Colmax - Collow + 1)] for i in range(Rowmax - Rowlow + 1)]
    for row in sheet.iter_rows():
        n = 0
        j = 0
        for cell in row:

            if m + 1 >= Rowlow and m + 1 <= Rowmax and n + 1 >= Collow and n + 1 <= Colmax:
                # 获取单元格的值
                a = cell.value
                SheetData[i][j] = a
                j = j + 1
            n = n + 1
        if m + 1 >= Rowlow and m + 1 <= Rowmax:
            i = i + 1
        m = m + 1
    if RowNum == 1:
        SheetSingle = {}
        i = 0
        for data in SheetData[0]:
            SheetSingle[i] = data
            i = i + 1
        SheetData = SheetSingle
    elif ColNum == 1:
        SheetSingle = {}
        for i in range(0, RowNum):
            SheetSingle[i] = SheetData[i][0]
        SheetData = SheetSingle
    return SheetData


def ExcleWrite(ExclePath, SheetIndex, Row, Col, Value, SaveAsNewFile):
    if SaveAsNewFile:
        FileNewName = "New_" + FileName(ExclePath)
        if FileFatherPath(ExclePath) != "":
            ExcleNewPath = FileFatherPath(ExclePath) + "\\" + FileNewName
        else:
            ExcleNewPath = FileNewName
        workbook = openpyxl.load_workbook(ExclePath)
        sheet_names = workbook.sheetnames
        SheetName = sheet_names[SheetIndex - 1]
        # 选择要操作的工作表
        sheet = workbook[SheetName]
        # 在指定的单元格写入数据
        sheet.cell(row=Row, column=Col, value=Value)
        # 保存文件
        workbook.save(ExcleNewPath)
    else:
        workbook = openpyxl.load_workbook(ExclePath)
        sheet_names = workbook.sheetnames
        SheetName = sheet_names[SheetIndex - 1]
        # 选择要操作的工作表
        sheet = workbook[SheetName]
        # 在指定的单元格写入数据
        sheet.cell(row=Row, column=Col, value=Value)
        # 保存文件
        workbook.save(ExclePath)


def WordTableRead(WordPath, TableIndex):
    doc = docx.Document(WordPath)
    table = doc.tables[TableIndex - 1]
    RowNum = 0
    for row in table.rows:
        ColNum = 0
        for cell in row.cells:
            ColNum = ColNum + 1
        RowNum = RowNum + 1
    SheetData = [[None for j in range(ColNum)] for i in range(RowNum)]
    i = 0

    for row in table.rows:
        j = 0
        for cell in row.cells:
            if i == 0 and j == 0:
                bcell_text = None
                cell_text = cell.text
                SheetData[i][j] = cell_text
            else:
                bcell_text = cell_text
                cell_text = cell.text
                if bcell_text != cell_text:
                    SheetData[i][j] = cell_text
                else:
                    SheetData[i][j] = None
            j = j + 1
        i = i + 1
    return SheetData


def WordTableWrite(WordPath, TableIndex, Row, Col, Text, SaveAsNewFile):
    if SaveAsNewFile:
        FileNewName = "New_" + FileName(WordPath)
        if FileFatherPath(WordPath) != "":
            WordNewPath = FileFatherPath(WordPath) + "\\" + FileNewName
        else:
            WordNewPath = FileNewName
        doc = docx.Document(WordPath)
        table = doc.tables[TableIndex - 1]
        Cell = table.cell(Row - 1, Col - 1)
        paragraph = Cell.paragraphs[0]  # 获取单元格中的段落
        alignmentOrl = paragraph.paragraph_format.alignment  # 获取单元格段落对齐状态
        ## 新增代码段 ##
        SuperTextBefor = ""
        SuperText = ""
        SuperTextAfter = ""
        SubTextBefor = ""
        SubText = ""
        SubTextAfter = ""
        for i in range(0, len(Text)):
            if Text[i] == "^" and i + 1 < len(Text):
                SubTextBefor = ""
                if Text[i + 1] == "(" and i + 2 < len(Text):
                    for j in range(i + 2, len(Text)):
                        if Text[j] == ")":
                            for k in range(j + 1, len(Text)):
                                SuperTextAfter = SuperTextAfter + Text[k]
                            break
                        SuperText = SuperText + Text[j]
                    break
                else:
                    for j in range(i + 1, len(Text)):
                        SuperText = SuperText + Text[j]
                    break
            elif Text[i] == "_" and i + 1 < len(Text):
                SuperTextBefor = ""
                if Text[i + 1] == "(" and i + 2 < len(Text):
                    for j in range(i + 2, len(Text)):
                        if Text[j] == ")":
                            for k in range(j + 1, len(Text)):
                                SubTextAfter = SubTextAfter + Text[k]
                            break
                        SubText = SubText + Text[j]
                    break
                else:
                    for j in range(i + 1, len(Text)):
                        SubText = SubText + Text[j]
                    break
            SuperTextBefor = SuperTextBefor + Text[i]
            SubTextBefor = SubTextBefor + Text[i]

        if SuperTextBefor == Text and SubTextBefor == Text:
            Cell.text = Text
        elif SuperText != "":
            Cell.text = SuperTextBefor
            paragraph = Cell.paragraphs[0]  # 获取单元格中的段落
            AddedSuperText = paragraph.add_run(SuperText)
            AddedSuperText.font.superscript = True
            paragraph.add_run(SuperTextAfter)
        else:
            Cell.text = SubTextBefor
            paragraph = Cell.paragraphs[0]  # 获取单元格中的段落
            AddedSuperText = paragraph.add_run(SubText)
            AddedSuperText.font.subscript = True
            paragraph.add_run(SubTextAfter)
        ## 新增代码段 ##
        for paragraph in Cell.paragraphs:
            paragraph.paragraph_format.left_indent = 0  # 预先对缩进赋值，防止对象为空报错
            paragraph.paragraph_format.element.pPr.ind.set(qn("w:firstLineChars"), '-1')  # 去除缩进
            paragraph_format = paragraph.paragraph_format
            paragraph_format.alignment = alignmentOrl  # 设置为原段落的对齐方式
        doc.save(WordNewPath)
    else:
        doc = docx.Document(WordPath)
        table = doc.tables[TableIndex - 1]
        Cell = table.cell(Row - 1, Col - 1)
        paragraph = Cell.paragraphs[0]  # 获取单元格中的段落
        alignmentOrl = paragraph.paragraph_format.alignment  # 获取单元格段落对齐状态
        ## 新增代码段 ##
        SuperTextBefor = ""
        SuperText = ""
        SuperTextAfter = ""
        SubTextBefor = ""
        SubText = ""
        SubTextAfter = ""
        for i in range(0, len(Text)):
            if Text[i] == "^" and i + 1 < len(Text):
                SubTextBefor = ""
                if Text[i + 1] == "(" and i + 2 < len(Text):
                    for j in range(i + 2, len(Text)):
                        if Text[j] == ")":
                            for k in range(j + 1, len(Text)):
                                SuperTextAfter = SuperTextAfter + Text[k]
                            break
                        SuperText = SuperText + Text[j]
                    break
                else:
                    for j in range(i + 1, len(Text)):
                        SuperText = SuperText + Text[j]
                    break
            elif Text[i] == "_" and i + 1 < len(Text):
                SuperTextBefor = ""
                if Text[i + 1] == "(" and i + 2 < len(Text):
                    for j in range(i + 2, len(Text)):
                        if Text[j] == ")":
                            for k in range(j + 1, len(Text)):
                                SubTextAfter = SubTextAfter + Text[k]
                            break
                        SubText = SubText + Text[j]
                    break
                else:
                    for j in range(i + 1, len(Text)):
                        SubText = SubText + Text[j]
                    break
            SuperTextBefor = SuperTextBefor + Text[i]
            SubTextBefor = SubTextBefor + Text[i]

        if SuperTextBefor == Text and SubTextBefor == Text:
            Cell.text = Text
        elif SuperText != "":
            Cell.text = SuperTextBefor
            paragraph = Cell.paragraphs[0]  # 获取单元格中的段落
            AddedSuperText = paragraph.add_run(SuperText)
            AddedSuperText.font.superscript = True
            paragraph.add_run(SuperTextAfter)
        else:
            Cell.text = SubTextBefor
            paragraph = Cell.paragraphs[0]  # 获取单元格中的段落
            AddedSuperText = paragraph.add_run(SubText)
            AddedSuperText.font.subscript = True
            paragraph.add_run(SubTextAfter)
        ## 新增代码段 ##
        for paragraph in Cell.paragraphs:
            paragraph.paragraph_format.left_indent = 0  # 预先对缩进赋值，防止对象为空报错
            paragraph.paragraph_format.element.pPr.ind.set(qn("w:firstLineChars"), '-1')  # 去除缩进
            paragraph_format = paragraph.paragraph_format
            paragraph_format.alignment = alignmentOrl  # 设置为原段落的对齐方式
        doc.save(WordPath)


def WordTableInsertFig(WordPath, TableIndex, Row, Col, ImagePath, ImageHeight_cm, ImageWidth_cm, SaveAsNewFile):
    if SaveAsNewFile:
        FileNewName = "New_" + FileName(WordPath)
        if FileFatherPath(WordPath) != "":
            WordNewPath = FileFatherPath(WordPath) + "\\" + FileNewName
        else:
            WordNewPath = FileNewName
        doc = docx.Document(WordPath)
        table = doc.tables[TableIndex - 1]
        Cell = table.cell(Row - 1, Col - 1)
        paragraph = Cell.paragraphs[0]  # 获取单元格中的段落
        alignmentOrl = paragraph.paragraph_format.alignment  # 获取单元格段落对齐状态
        if ImageHeight_cm is not None and ImageWidth_cm is not None:
            Cell.add_paragraph().add_run().add_picture(ImagePath, height=Cm(ImageHeight_cm), width=Cm(ImageWidth_cm))
        elif ImageHeight_cm is not None:
            Cell.add_paragraph().add_run().add_picture(ImagePath, height=Cm(ImageHeight_cm))
        elif ImageWidth_cm is not None:
            Cell.add_paragraph().add_run().add_picture(ImagePath, width=Cm(ImageWidth_cm))
        else:
            Cell.add_paragraph().add_run().add_picture(ImagePath)
        for paragraph in Cell.paragraphs:
            paragraph.paragraph_format.left_indent = 0  # 预先对缩进赋值，防止对象为空报错
            paragraph.paragraph_format.element.pPr.ind.set(qn("w:firstLineChars"), '-1')  # 去除缩进
            paragraph_format = paragraph.paragraph_format
            paragraph_format.alignment = alignmentOrl  # 设置为原段落的对齐方式
        doc.save(WordNewPath)
    else:
        doc = docx.Document(WordPath)
        table = doc.tables[TableIndex - 1]
        Cell = table.cell(Row - 1, Col - 1)
        paragraph = Cell.paragraphs[0]  # 获取单元格中的段落
        alignmentOrl = paragraph.paragraph_format.alignment  # 获取单元格段落对齐状态
        if ImageHeight_cm is not None and ImageWidth_cm is not None:
            Cell.add_paragraph().add_run().add_picture(ImagePath, height=Cm(ImageHeight_cm), width=Cm(ImageWidth_cm))
        elif ImageHeight_cm is not None:
            Cell.add_paragraph().add_run().add_picture(ImagePath, height=Cm(ImageHeight_cm))
        elif ImageWidth_cm is not None:
            Cell.add_paragraph().add_run().add_picture(ImagePath, width=Cm(ImageWidth_cm))
        else:
            Cell.add_paragraph().add_run().add_picture(ImagePath)
        for paragraph in Cell.paragraphs:
            paragraph.paragraph_format.left_indent = 0  # 预先对缩进赋值，防止对象为空报错
            paragraph.paragraph_format.element.pPr.ind.set(qn("w:firstLineChars"), '-1')  # 去除缩进
            paragraph_format = paragraph.paragraph_format
            paragraph_format.alignment = alignmentOrl  # 设置为原段落的对齐方式
        doc.save(WordPath)


def WordTableParaAlignment(WordPath, TableIndex, Row, Col, Alignment_left_right_center_None, SaveAsNewFile):
    if SaveAsNewFile:
        FileNewName = "New_" + FileName(WordPath)
        if FileFatherPath(WordPath) != "":
            WordNewPath = FileFatherPath(WordPath) + "\\" + FileNewName
        else:
            WordNewPath = FileNewName
        doc = docx.Document(WordPath)
        table = doc.tables[TableIndex - 1]
        Cell = table.cell(Row - 1, Col - 1)
        if Alignment_left_right_center_None == "left":
            alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif Alignment_left_right_center_None == "right":
            alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        elif Alignment_left_right_center_None == "center":
            alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        else:
            alignment = None
        for paragraph in Cell.paragraphs:
            paragraph.paragraph_format.left_indent = 0  # 预先对缩进赋值，防止对象为空报错
            paragraph.paragraph_format.element.pPr.ind.set(qn("w:firstLineChars"), '-1')  # 去除缩进
            paragraph_format = paragraph.paragraph_format
            paragraph_format.alignment = alignment  # 设置为原段落的对齐方式
        doc.save(WordNewPath)
    else:
        doc = docx.Document(WordPath)
        table = doc.tables[TableIndex - 1]
        Cell = table.cell(Row - 1, Col - 1)
        if Alignment_left_right_center_None == "left":
            alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif Alignment_left_right_center_None == "right":
            alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        elif Alignment_left_right_center_None == "center":
            alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        else:
            alignment = None
        for paragraph in Cell.paragraphs:
            paragraph.paragraph_format.left_indent = 0  # 预先对缩进赋值，防止对象为空报错
            paragraph.paragraph_format.element.pPr.ind.set(qn("w:firstLineChars"), '-1')  # 去除缩进
            paragraph_format = paragraph.paragraph_format
            paragraph_format.alignment = alignment  # 设置为原段落的对齐方式
        doc.save(WordPath)


def ShpToXlsx(ShpPath, XlsxPath):
    with fiona.open(ShpPath, "r") as shapefile:
        # 读取SHP文件的属性表
        properties = [feature["properties"] for feature in shapefile]

        # 将属性表转换为DataFrame
        df = pd.DataFrame(properties)

        # 将DataFrame写入Excel文件
        df.to_excel(XlsxPath, index=False)


def wordAdd(wordPath, wordSavePath, new_text, FontName=None, FontSize=None, IsBold=None, IsItalic=None):
    doc = Document(wordPath)
    # 获取文档的最后一段
    last_paragraph = doc.paragraphs[-1]

    if last_paragraph.runs:
        # 获取最后一段的最后一个run
        last_run = last_paragraph.runs[-1]
        # 创建一个新的run
        new_run = last_paragraph.add_run(new_text)

        # 复制字体名称
        if FontName is None:
            if last_run.font.name:
                new_run.font.name = last_run.font.name
                new_run.font._element.rPr.rFonts.set(qn('w:eastAsia'), last_run.font.name)
        else:
            new_run.font.name = FontName
            new_run.font._element.rPr.rFonts.set(qn('w:eastAsia'), FontName)
        # 复制字体大小
        if FontSize is None:
            if last_run.font.size:
                new_run.font.size = last_run.font.size
        else:
            new_run.font.size = FontSize
        # 复制加粗、斜体等其他属性
        if IsBold is None:
            new_run.bold = last_run.bold
        else:
            if IsBold:
                new_run.bold = True
            else:
                new_run.bold = False
        if IsItalic is None:
            new_run.italic = last_run.italic
        else:
            if IsItalic:
                new_run.italic = True
            else:
                new_run.italic = False
    else:
        # 如果最后一段是空的，则直接添加新文本
        last_paragraph.add_run(new_text)

    # 保存修改后的文档
    doc.save(wordSavePath)


def wordParagraphAdd(wordPath, wordSavePath, new_text, FontName=None, FontSize=None, IsBold=None, IsItalic=None,
                     Indent=None, Alignment="l"):
    doc = Document(wordPath)
    # 获取文档的最后一段
    last_paragraph = doc.paragraphs[-1]

    if last_paragraph.runs:
        # 获取最后一段的最后一个run
        last_run = last_paragraph.runs[-1]
        # 创建一个新的run
        new_paragraph = doc.add_paragraph()
        new_run = new_paragraph.add_run(new_text)

        format = last_paragraph.paragraph_format

        # 获取左缩进和首行缩进（单位是英寸，Word中的默认单位）
        left_indent = format.left_indent
        first_line_indent = format.first_line_indent
        if Indent is None:
            new_paragraph.paragraph_format.first_line_indent = first_line_indent
        else:
            try:
                new_paragraph.paragraph_format.first_line_indent = Indent
            except:
                new_paragraph.paragraph_format.first_line_indent = first_line_indent
        if Alignment == "l":
            new_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif Alignment == "c":
            new_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif Alignment == "r":
            new_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif Alignment == "j":
            new_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            new_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # 复制字体名称
        if FontName is None:
            if last_run.font.name:
                new_run.font.name = last_run.font.name
                new_run.font._element.rPr.rFonts.set(qn('w:eastAsia'), last_run.font.name)
        else:
            new_run.font.name = FontName
            new_run.font._element.rPr.rFonts.set(qn('w:eastAsia'), FontName)
        # 复制字体大小
        if FontSize is None:
            if last_run.font.size:
                new_run.font.size = last_run.font.size
        else:
            new_run.font.size = FontSize
        # 复制加粗、斜体等其他属性
        if IsBold is None:
            new_run.bold = last_run.bold
        else:
            if IsBold:
                new_run.bold = True
            else:
                new_run.bold = False
        if IsItalic is None:
            new_run.italic = last_run.italic
        else:
            if IsItalic:
                new_run.italic = True
            else:
                new_run.italic = False
    else:
        # 如果最后一段是空的，则直接添加新文本
        last_paragraph.add_run(new_text)

    # 保存修改后的文档
    doc.save(wordSavePath)


def wordParaFormat(wordPath):
    doc = Document(wordPath)
    # 获取文档的最后一段
    last_paragraph = doc.paragraphs[-1]
    last_run = last_paragraph.runs[-1]
    format = last_paragraph.paragraph_format
    return format
